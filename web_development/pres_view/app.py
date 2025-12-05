from flask import (
    Blueprint,
    render_template,
    request,
    redirect,
    url_for,
    flash,
    session,
    jsonify,
    send_file,
)
import os
import re
from uuid import uuid4
from dotenv import load_dotenv
from werkzeug.security import check_password_hash, generate_password_hash
from supabase import create_client
from slugify import slugify
from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import datetime as dt

pres = Blueprint(
    "pres",
    __name__,
    template_folder="templates/pres",
    static_folder="static",
)

load_dotenv()

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

BUCKET_RECEIPTS = "Receipts"


# -----------------------
# Helpers
# -----------------------


def validate_password(pw: str):
    errors = []
    if len(pw) < 8:
        errors.append("Password must be at least 8 characters.")
    if not re.search(r"[a-z]", pw):
        errors.append("Password must contain a lowercase letter.")
    if not re.search(r"[A-Z]", pw):
        errors.append("Password must contain an uppercase letter.")
    if not re.search(r"\d", pw):
        errors.append("Password must contain a number.")
    if not re.search(r"[^A-Za-z0-9]", pw):
        errors.append("Password must contain a special character.")
    return errors


def create_osas_notification(org_id, report_id, message=None):
    """
    Insert a row into osas_notifications so OSAS bell sees a new item.
    """
    try:
        # kunin org_name para sa display
        org_res = (
            supabase.table("organizations")
            .select("org_name")
            .eq("id", org_id)
            .single()
            .execute()
        )
        org_name = org_res.data["org_name"] if org_res.data else "Organization"

        if not message:
            message = 'has a report "Pending Review"'

        supabase.table("osas_notifications").insert(
            {
                "org_id": org_id,
                "report_id": report_id,
                "org_name": org_name,
                "message": message,
                # created_at at is_read may default na sa table definition
            }
        ).execute()
    except Exception:
        # huwag pabagsakin ang submit kahit mag-fail ang notif
        pass


def get_real_wallet_id(folder_id: int):
    """
    folder_id = wallet_budgets.id â†’ return wallet_id or None.
    """
    res = (
        supabase.table("wallet_budgets")
        .select("wallet_id")
        .eq("id", folder_id)
        .single()
        .execute()
    )
    return res.data["wallet_id"] if res.data else None


# -----------------------
# Landing + health
# -----------------------


@pres.route("/")
def landingpage():
    if request.accept_mimetypes.best == "application/json":
        return jsonify({"status": "ok", "page": "landing"})
    return render_template("landingpage.html")


@pres.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "api": "PRES API", "version": "1.0"})


# -----------------------
# Login / auth
# -----------------------


@pres.route("/login/pres", methods=["GET", "POST"])
def pres_login():
    if request.method == "POST":
        username = request.form.get("username") or (
            request.json.get("username") if request.is_json else None
        )
        password = request.form.get("password") or (
            request.json.get("password") if request.is_json else None
        )

        result = (
            supabase.table("organizations")
            .select("*")
            .eq("username", username)
            .execute()
        )

        if result.data:
            org = result.data[0]

            # block Archived orgs
            if org.get("status") == "Archived":
                error_msg = (
                    "This organization account is archived. Please contact OSAS."
                )
                if request.accept_mimetypes.best == "application/json":
                    return jsonify({"success": False, "error": error_msg}), 403
                flash(error_msg, "danger")
                return redirect(url_for("pres.pres_login"))

            if check_password_hash(org["password"], password):
                session["pres_user"] = True
                session["org_id"] = org["id"]
                session["org_name"] = org["org_name"]
                ...

                if request.accept_mimetypes.best == "application/json":
                    if org.get("must_change_password", False):
                        return jsonify(
                            {
                                "success": True,
                                "must_change_password": True,
                                "org_id": org["id"],
                                "org_name": org["org_name"],
                            }
                        )
                    return jsonify(
                        {
                            "success": True,
                            "org_id": org["id"],
                            "org_name": org["org_name"],
                        }
                    )

                if org.get("must_change_password", False):
                    return redirect(url_for("pres.change_password"))
                return redirect(url_for("pres.homepage"))
            else:
                error_msg = "Incorrect password."
                if request.accept_mimetypes.best == "application/json":
                    return jsonify({"success": False, "error": error_msg}), 401
                flash(error_msg, "danger")
                return redirect(url_for("pres.pres_login"))
        else:
            error_msg = "Organization not found."
            if request.accept_mimetypes.best == "application/json":
                return jsonify({"success": False, "error": error_msg}), 404
            flash(error_msg, "danger")
            return redirect(url_for("pres.pres_login"))

    if request.accept_mimetypes.best == "application/json":
        return jsonify({"info": "POST username and password to this endpoint."})
    return render_template("pres_login.html")


@pres.route("/api/auth_status", methods=["GET"])
def pres_auth_status():
    login_state = session.get("pres_user") is True
    return jsonify(
        {
            "loggedin": login_state,  # ← camelCase!
            "orgid": session.get("org_id") if login_state else None,  # ← camelCase!
            "org_name": session.get("org_name") if login_state else None,
        }
    )


# -----------------------
# Forgot / change password
# -----------------------


@pres.route("/forgot-password", methods=["GET", "POST"])
def pres_forgot_password():
    if request.method == "POST":
        username = request.form.get("username")
        # TODO: real reset flow
        flash("Password reset instructions sent!", "success")
    return render_template("forgot_password.html")


@pres.route("/change-password", methods=["GET", "POST"])
def change_password():
    if not session.get("pres_user"):
        return redirect(url_for("pres.pres_login"))
    if "org_id" not in session:
        flash("Session expired. Please log in again.", "danger")
        return redirect(url_for("pres.pres_login"))

    if request.method == "POST":
        new_pw = request.form.get("new_password") or (
            request.json.get("new_password") if request.is_json else None
        )
        confirm_pw = request.form.get("confirm_password") or (
            request.json.get("confirm_password") if request.is_json else None
        )

        if not new_pw or not confirm_pw:
            error_msg = "Please fill out all fields."
            if request.accept_mimetypes.best == "application/json":
                return jsonify({"success": False, "error": error_msg}), 400
            flash(error_msg, "danger")
            return render_template("change_password.html")

        if new_pw != confirm_pw:
            error_msg = "Passwords do not match."
            if request.accept_mimetypes.best == "application/json":
                return jsonify({"success": False, "error": error_msg}), 400
            flash(error_msg, "danger")
            return render_template("change_password.html")

        errors = validate_password(new_pw)
        if errors:
            if request.accept_mimetypes.best == "application/json":
                return jsonify({"success": False, "errors": errors}), 400
            for e in errors:
                flash(e, "danger")
            return render_template("change_password.html")

        org_id = session.get("org_id")
        hashed_pw = generate_password_hash(new_pw)
        supabase.table("organizations").update(
            {"password": hashed_pw, "must_change_password": False}
        ).eq("id", org_id).execute()
        session["pres_user"] = True

        if request.accept_mimetypes.best == "application/json":
            return jsonify({"success": True, "message": "Password changed"})
        flash("Password changed successfully!", "success")
        return redirect(url_for("pres.homepage"))

    if request.accept_mimetypes.best == "application/json":
        return jsonify({"info": "POST new_password and confirm_password"})
    return render_template("change_password.html")


# -----------------------
# Basic pages
# -----------------------


@pres.route("/homepage")
def homepage():
    if not session.get("pres_user"):
        return redirect(url_for("pres.pres_login"))
    return render_template("pres_homepage.html")


@pres.route("/history")
def history():
    if not session.get("pres_user"):
        return redirect(url_for("pres.pres_login"))
    return render_template("history.html")


@pres.route("/wallets")
def wallets():
    if not session.get("pres_user"):
        return redirect(url_for("pres.pres_login"))
    return render_template("wallets.html")


@pres.route("/profile")
def profile():
    if not session.get("pres_user"):
        return redirect(url_for("pres.pres_login"))
    return render_template("profile.html")


@pres.route("/logout")
def pres_logout():
    session.pop("pres_user", None)
    session.pop("org_id", None)
    session.pop("org_name", None)
    return redirect(url_for("pres.pres_login"))


# -----------------------
# Dashboard summary (dummy)
# -----------------------


@pres.route("/api/dashboard/summary", methods=["GET"])
def get_dashboard_summary():
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")

    from datetime import datetime

    try:
        # 1) get all wallets for this org
        wallets_res = (
            supabase.table("wallets")
            .select("id")
            .eq("organization_id", org_id)
            .execute()
        )
        wallet_ids = [w["id"] for w in (wallets_res.data or [])]
        if not wallet_ids:
            return jsonify(
                {
                    "total_balance": 0,
                    "reports_submitted": 0,
                    "income_month": 0,
                    "expenses_month": 0,
                }
            )

        # 2) get all transactions for those wallets
        tx_res = (
            supabase.table("wallet_transactions")
            .select("kind, date_issued, quantity, price")
            .in_("wallet_id", wallet_ids)
            .execute()
        )

        now = datetime.now()
        this_year = now.year
        this_month = now.month

        income_month = 0.0
        expenses_month = 0.0
        total_income_all = 0.0
        total_expenses_all = 0.0

        for tx in (tx_res.data or []):
            qty = int(tx.get("quantity") or 0)
            price = float(tx.get("price") or 0)
            amt = qty * price

            # total across all time for balance
            if tx.get("kind") == "income":
                total_income_all += amt
            elif tx.get("kind") == "expense":
                total_expenses_all += amt

            # filter by current month/year for monthly cards
            d_str = tx.get("date_issued")
            try:
                d = datetime.fromisoformat(d_str.replace("Z", "+00:00"))
            except Exception:
                continue

            if d.year == this_year and d.month == this_month:
                if tx.get("kind") == "income":
                    income_month += amt
                elif tx.get("kind") == "expense":
                    expenses_month += amt

        # 3) budgets (beginning cash) to compute total balance
        budgets_res = (
            supabase.table("wallet_budgets")
            .select("amount, wallet_id")
            .in_("wallet_id", wallet_ids)
            .execute()
        )
        beginning_total = sum(
            float(b.get("amount") or 0) for b in (budgets_res.data or [])
        )

        total_balance = beginning_total + total_income_all - total_expenses_all

        # 4) reports submitted (any financial_reports for this org)
        reports_res = (
            supabase.table("financial_reports")
            .select("id", count="exact")
            .eq("organization_id", org_id)
            .in_("status", ["Submitted", "Approved"])
            .execute()
        )
        reports_submitted = reports_res.count or 0

        return jsonify(
            {
                "total_balance": total_balance,
                "reports_submitted": reports_submitted,
                "income_month": income_month,
                "expenses_month": expenses_month,
            }
        )
    except Exception as e:
        print("Error get_dashboard_summary:", e)
        return jsonify(
            {
                "total_balance": 0,
                "reports_submitted": 0,
                "income_month": 0,
                "expenses_month": 0,
                "error": str(e),
            }
        ), 500


# -----------------------
# API: wallets -> month folders
# -----------------------

@pres.route("/api/wallets", methods=["GET"])
def get_wallets():
    """Return wallets with exact field names matching JS"""
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")

    try:
        # Get all wallets for org
        wallets_res = (
            supabase.table("wallets")
            .select("id, name")
            .eq("organization_id", org_id)
            .execute()
        )

        if not wallets_res.data:
            return jsonify([])

        result = []
        for w in wallets_res.data:
            wallet_id = w["id"]

            # Get budgets for this wallet (with year + month info)
            budgets_res = (
                supabase.table("wallet_budgets")
                .select("id, amount, year, month_id, months(month_name, month_order)")
                .eq("wallet_id", wallet_id)
                .order("year")
                .order("month_id")
                .execute()
            )

            for row in budgets_res.data or []:
                m = row["months"]
                year = row["year"]
                month_order = m["month_order"]
                month_code = f"{year}-{month_order:02d}"

                # ✅ EXACT field names the JS expects!
                result.append(
                    {
                        "id": row["id"],
                        "wallet_id": wallet_id,
                        "name": m["month_name"],
                        "month": month_code,
                        "beginning_cash": float(row["amount"] or 0),
                    }
                )

        return jsonify(result)
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": str(e)}), 500
    
@pres.route("/api/wallets/overview", methods=["GET"])
def get_wallets_overview():
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")

    try:
        # 1) all wallets for this org
        wallets_res = (
            supabase.table("wallets")
            .select("id, name")
            .eq("organization_id", org_id)
            .execute()
        )
        wallets = {w["id"]: w["name"] for w in (wallets_res.data or [])}
        if not wallets:
            return jsonify([])

        # 2) all budget folders for those wallets
        budgets_res = (
            supabase.table("wallet_budgets")
            .select("id, wallet_id, amount, year, month_id, months (month_name)")
            .in_("wallet_id", list(wallets.keys()))
            .execute()
        )
        budget_ids = [b["id"] for b in (budgets_res.data or [])]
        if not budget_ids:
            return jsonify([])

        # 3) all transactions for those folders
        tx_res = (
            supabase.table("wallet_transactions")
            .select("budget_id, kind, quantity, price, date_issued")
            .in_("budget_id", budget_ids)
            .execute()
        )

        from datetime import datetime

        # init per-folder aggregates
        by_folder = {}
        for b in budgets_res.data or []:
            fid = b["id"]
            by_folder[fid] = {
                "id": fid,
                "wallet_id": b["wallet_id"],
                "name": b["months"]["month_name"],          # title, e.g. DECEMBER
                "budget": float(b.get("amount") or 0),      # budget for that month
                "total_income": 0.0,
                "total_expenses": 0.0,
                "last_activity": None,
            }

        # aggregate income/expenses by folder
        for tx in tx_res.data or []:
            fid = tx["budget_id"]
            if fid not in by_folder:
                continue
            qty = int(tx.get("quantity") or 0)
            price = float(tx.get("price") or 0)
            amt = qty * price
            if tx.get("kind") == "income":
                by_folder[fid]["total_income"] += amt
            elif tx.get("kind") == "expense":
                by_folder[fid]["total_expenses"] += amt

            ts = tx.get("date_issued")
            if ts:
                try:
                    dt_val = datetime.fromisoformat(ts.replace("Z", "+00:00"))
                except Exception:
                    dt_val = None
                if dt_val:
                    cur = by_folder[fid]["last_activity"]
                    if not cur or dt_val > cur:
                        by_folder[fid]["last_activity"] = dt_val

        # keep only folders that have any transaction
        items = [
            v for v in by_folder.values()
            if v["total_income"] > 0 or v["total_expenses"] > 0
        ]

        # sort so recently edited/added are on top
        items.sort(
            key=lambda x: x["last_activity"] or datetime.min,
            reverse=True,
        )

        return jsonify(items)
    except Exception as e:
        print("Error get_wallets_overview:", e)
        return jsonify({"error": str(e)}), 500




# ----------------------
# Budget endpoints (per folder)
# -----------------------


@pres.route("/api/wallets/<int:folder_id>/budget/current-month", methods=["GET"])
def get_budget(folder_id):
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401
    try:
        res = (
            supabase.table("wallet_budgets")
            .select("amount")
            .eq("id", folder_id)
            .single()
            .execute()
        )
        if not res.data:
            return jsonify({"amount": None})
        return jsonify({"amount": res.data["amount"]})
    except Exception as e:
        return jsonify({"amount": None, "error": str(e)}), 500


@pres.route("/api/wallets/<int:folder_id>/budget", methods=["POST"])
def set_budget(folder_id):
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    data = request.get_json() or {}
    amount = data.get("amount", 0)

    try:
        supabase.table("wallet_budgets").update({"amount": amount}).eq(
            "id", folder_id
        ).execute()
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# -----------------------
# Wallet transactions (per folder)
# -----------------------


@pres.route("/api/wallets/<int:folder_id>/transactions", methods=["GET"])
def get_wallet_transactions(folder_id):
    """Get transactions for wallet"""
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    try:
        # Get folder to find wallet_id
        folder_res = (
            supabase.table("wallet_budgets")
            .select("wallet_id")
            .eq("id", folder_id)
            .execute()
        )

        if not folder_res.data:
            return jsonify([])

        wallet_id = folder_res.data[0]["wallet_id"]

        res = (
            supabase.table("wallet_transactions")
            .select(
                "id, kind, date_issued, description, quantity, price, "
                "income_type, particulars"
            )
            .eq("wallet_id", wallet_id)
            .eq("budget_id", folder_id)
            .order("date_issued")
            .execute()
        )

        txs = []
        for tx in res.data or []:
            qty = int(tx.get("quantity", 0))
            price = float(tx.get("price", 0))
            total_amount = qty * price

            txs.append(
                {
                    "id": tx["id"],
                    "quantity": qty,
                    "price": price,
                    "incometype": tx.get("income_type"),
                    "particulars": tx.get("particulars"),
                    "description": tx["description"],
                    "total_amount": total_amount,
                    "date_issued": tx["date_issued"],
                    "kind": tx["kind"],
                }
            )
        return jsonify(txs)
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": str(e)}), 500


@pres.route("/api/wallets/<int:folder_id>/transactions", methods=["POST"])
def add_wallet_transaction(folder_id):
    """Create new transaction"""
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    wallet_id = get_real_wallet_id(folder_id)
    if wallet_id is None:
        return jsonify({"error": "Wallet folder not found"}), 404

    data = request.get_json() or {}
    kind = data.get("kind")
    date_issued = data.get("date_issued")
    quantity = data.get("quantity")
    income_type = data.get("income_type")
    particulars = data.get("particulars")
    description = data.get("description")
    price = data.get("price")

    if kind not in ("income", "expense"):
        return jsonify({"error": "Invalid kind"}), 400

    try:
        ins = (
            supabase.table("wallet_transactions")
            .insert(
                {
                    "wallet_id": wallet_id,
                    "budget_id": folder_id,
                    "kind": kind,
                    "date_issued": date_issued,
                    "quantity": quantity,
                    "income_type": income_type,
                    "particulars": particulars,
                    "description": description,
                    "price": price,
                }
            )
            .execute()
        )
        row = ins.data[0]
        total_amount = float(row["price"]) * int(row["quantity"])

        return jsonify(
            {
                "transaction": {
                    "id": row["id"],
                    "wallet_id": row["wallet_id"],
                    "kind": row["kind"],
                    "date_issued": row["date_issued"],
                    "description": row["description"],
                    "quantity": row["quantity"],
                    "price": float(row["price"]),
                    "total_amount": total_amount,
                    "income_type": row.get("income_type"),
                    "particulars": row.get("particulars"),
                }
            }
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@pres.route("/api/wallets/<int:folder_id>/transactions/<int:tx_id>", methods=["POST"])
def update_wallet_transaction(folder_id, tx_id):
    """Update existing transaction"""
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    wallet_id = get_real_wallet_id(folder_id)
    if wallet_id is None:
        return jsonify({"error": "Wallet folder not found"}), 404

    data = request.get_json() or {}
    kind = data.get("kind")
    date_issued = data.get("date_issued")
    quantity = data.get("quantity")
    income_type = data.get("income_type")
    particulars = data.get("particulars")
    description = data.get("description")
    price = data.get("price")

    if kind not in ("income", "expense"):
        return jsonify({"error": "Invalid kind"}), 400

    try:
        upd = (
            supabase.table("wallet_transactions")
            .update(
                {
                    "kind": kind,
                    "date_issued": date_issued,
                    "quantity": quantity,
                    "income_type": income_type,
                    "particulars": particulars,
                    "description": description,
                    "price": price,
                }
            )
            .eq("id", tx_id)
            .execute()
        )
        row = upd.data[0]
        total_amount = float(row["price"]) * int(row["quantity"])

        return jsonify(
            {
                "transaction": {
                    "id": row["id"],
                    "wallet_id": row["wallet_id"],
                    "kind": row["kind"],
                    "date_issued": row["date_issued"],
                    "description": row["description"],
                    "quantity": row["quantity"],
                    "price": float(row["price"]),
                    "total_amount": total_amount,
                    "income_type": row.get("income_type"),
                    "particulars": row.get("particulars"),
                }
            }
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@pres.route("/api/wallets/<int:folder_id>/transactions/<int:tx_id>", methods=["DELETE"])
def delete_wallet_transaction(folder_id, tx_id):
    """Delete transaction"""
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    try:
        supabase.table("wallet_transactions").delete().eq("id", tx_id).execute()
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
@pres.route("/api/transactions/recent", methods=["GET"])
def get_recent_transactions():
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")

    try:
        # get all wallets of this org
        wallets_res = (
            supabase.table("wallets")
            .select("id")
            .eq("organization_id", org_id)
            .execute()
        )
        wallet_ids = [w["id"] for w in (wallets_res.data or [])]
        if not wallet_ids:
            return jsonify([])

        res = (
            supabase.table("wallet_transactions")
            .select("id, kind, date_issued, description, quantity, price, particulars")
            .in_("wallet_id", wallet_ids)
            .order("date_issued", desc=True)
            .limit(50)
            .execute()
        )

        txs = []
        for tx in res.data or []:
            qty = int(tx.get("quantity") or 0)
            price = float(tx.get("price") or 0)
            total_amount = qty * price

            txs.append(
                {
                    "id": tx["id"],
                    "type": "income" if tx.get("kind") == "income" else "expense",
                    "date": tx["date_issued"],
                    "event": tx.get("particulars") or "Transaction",
                    "description": tx.get("description") or "",
                    "amount": total_amount if tx.get("kind") == "income" else -total_amount,
                }
            )

        return jsonify(txs)
    except Exception as e:
        print("Error get_recent_transactions:", e)
        return jsonify({"error": str(e)}), 500




# -----------------------
# Wallet receipts + Storage (per folder)
# -----------------------


@pres.route("/api/wallets/<int:folder_id>/receipts", methods=["GET"])
def get_wallet_receipts(folder_id):
    """Get receipts for wallet"""
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    try:
        # Get folder to find wallet_id
        folder_res = (
            supabase.table("wallet_budgets")
            .select("wallet_id")
            .eq("id", folder_id)
            .execute()
        )

        if not folder_res.data:
            return jsonify([])

        wallet_id = folder_res.data[0]["wallet_id"]

        res = (
            supabase.table("wallet_receipts")
            .select("id, description, receipt_date, file_url")
            .eq("wallet_id", wallet_id)
            .eq("budget_id", folder_id)
            .order("receipt_date", desc=True)
            .execute()
        )

        receipts = [
            {
                "id": r["id"],
                "description": r["description"],
                "receipt_date": r["receipt_date"],
                "file_url": r["file_url"],
            }
            for r in (res.data or [])
        ]
        return jsonify(receipts)
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": str(e)}), 500


@pres.route("/api/wallets/<int:folder_id>/receipts", methods=["POST"])
def upload_wallet_receipt(folder_id):
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")
    wallet_id = get_real_wallet_id(folder_id)
    if wallet_id is None:
        return jsonify({"error": "Wallet folder not found"}), 404

    if "receipt-file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["receipt-file"]
    desc = request.form.get("receipt-desc")
    date = request.form.get("receipt-date")

    if not desc or not date:
        return jsonify({"error": "Missing fields"}), 400

    # org name
    org_res = (
        supabase.table("organizations")
        .select("org_name")
        .eq("id", org_id)
        .single()
        .execute()
    )
    org_name = org_res.data["org_name"] if org_res.data else f"org-{org_id}"

    # wallet name
    w_res = (
        supabase.table("wallets").select("name").eq("id", wallet_id).single().execute()
    )
    wallet_name = w_res.data["name"] if w_res.data else f"wallet-{wallet_id}"

    # month name
    b_res = (
        supabase.table("wallet_budgets")
        .select("month_id, months (month_name)")
        .eq("id", folder_id)
        .single()
        .execute()
    )
    month_name = b_res.data["months"]["month_name"] if b_res.data else "UNKNOWN"

    org_folder = slugify(org_name)
    wallet_folder = slugify(wallet_name)
    month_folder = slugify(month_name)

    ext = os.path.splitext(file.filename)[1] or ".png"
    safe_name = f"{uuid4()}{ext}"

    path = f"{org_folder}/{wallet_folder}/{month_folder}/{safe_name}"

    try:
        file_bytes = file.read()

        supabase.storage.from_(BUCKET_RECEIPTS).upload(path, file_bytes)

        ins = (
            supabase.table("wallet_receipts")
            .insert(
                {
                    "wallet_id": wallet_id,
                    "budget_id": folder_id,
                    "file_url": path,
                    "description": desc,
                    "receipt_date": date,
                }
            )
            .execute()
        )
        row = ins.data[0]

        return jsonify(
            {"id": row["id"], "name": row["description"], "date": row["receipt_date"]}
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@pres.route("/api/receipts/<int:receipt_id>/url", methods=["GET"])
def get_receipt_public_url(receipt_id):
    try:
        res = (
            supabase.table("wallet_receipts")
            .select("file_url")
            .eq("id", receipt_id)
            .single()
            .execute()
        )
        if not res.data:
            return jsonify({"error": "Not found"}), 404

        file_path = res.data["file_url"]
        public_url = supabase.storage.from_(BUCKET_RECEIPTS).get_public_url(file_path)

        if isinstance(public_url, dict):
            url = public_url.get("publicUrl") or public_url.get("signedURL")
        else:
            url = public_url

        return jsonify({"url": url})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@pres.route("/api/receipts/<int:receipt_id>/download-url", methods=["GET"])
def get_receipt_download_url(receipt_id):
    try:
        res = (
            supabase.table("wallet_receipts")
            .select("file_url")
            .eq("id", receipt_id)
            .single()
            .execute()
        )
        if not res.data:
            return jsonify({"error": "Not found"}), 404

        file_path = res.data["file_url"]
        public_url = supabase.storage.from_(BUCKET_RECEIPTS).get_public_url(
            file_path, {"download": True}
        )

        if isinstance(public_url, dict):
            url = public_url.get("publicUrl") or public_url.get("signedURL")
        else:
            url = public_url

        return jsonify({"url": url})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@pres.route("/api/receipts/<int:receipt_id>", methods=["DELETE"])
def delete_receipt(receipt_id):
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    try:
        res = (
            supabase.table("wallet_receipts")
            .select("file_url")
            .eq("id", receipt_id)
            .single()
            .execute()
        )
        if not res.data:
            return jsonify({"error": "Not found"}), 404

        file_path = res.data["file_url"]
        supabase.storage.from_(BUCKET_RECEIPTS).remove([file_path])
        supabase.table("wallet_receipts").delete().eq("id", receipt_id).execute()

        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# -----------------------
# Financial reports (PRES)
# -----------------------


@pres.route("/api/reports/generate", methods=["POST"])
def generate_report():
    """
    Create or update a single Pending report per wallet+budget.
    Called from PRES (user) side when they click Generate.
    """
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    data = request.get_json() or {}
    org_id = session.get("org_id")
    wallet_id = data.get("wallet_id")
    budget_id = data.get("budget_id")

    try:
        # Hanapin kung may existing Pending Review report para sa wallet+budget na ito
        existing = (
            supabase.table("financial_reports")
            .select("id")
            .eq("organization_id", org_id)
            .eq("wallet_id", wallet_id)
            .eq("budget_id", budget_id)
            .eq("status", "Pending Review")
            .limit(1)
            .execute()
        )

        # NEW: compute report_month galing sa wallet_budgets + months
        wb_res = (
            supabase.table("wallet_budgets")
            .select("year, month_id, months (month_name)")
            .eq("id", budget_id)
            .single()
            .execute()
        )
        report_month_value = None
        if wb_res.data:
            report_month_value = (wb_res.data["months"]["month_name"] or "").lower()

        # Fields dapat tugma sa JS payload at sa DOCX placeholders
        payload = {
            "organization_id": org_id,
            "wallet_id": wallet_id,
            "budget_id": budget_id,
            "status": "Pending Review",
            "notes": None,
            "checklist": {},
            "event_name": data.get("event_name"),
            "date_prepared": data.get("date_prepared"),
            "report_no": data.get("report_no"),
            "budget": data.get("budget"),
            "total_income": data.get("total_income"),  # {{TOTAL_INCOME}}
            "total_expense": data.get("total_expense"),
            "reimbursement": data.get("reimbursement"),
            "previous_fund": data.get("previous_fund"),
            "budget_in_the_bank": data.get(
                "budget_in_the_bank"
            ),  # {{BUDGET_IN_THE_BANK}}
            "report_month": report_month_value,
        }

        if existing.data:
            rep_id = existing.data[0]["id"]
            supabase.table("financial_reports").update(payload).eq(
                "id", rep_id
            ).execute()
        else:
            ins = supabase.table("financial_reports").insert(payload).execute()
            rep_id = ins.data[0]["id"]

        return jsonify({"success": True, "id": rep_id})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@pres.route("/api/wallets/reports/status", methods=["GET"])
def wallets_report_status():
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")
    month_code = request.args.get("month")  # e.g. "2025-12" or "december"
    if not month_code:
        return jsonify({"error": "Missing month"}), 400

    res = (
        supabase.table("financial_reports")
        .select("id, status, report_month")
        .eq("organization_id", org_id)
        .eq("report_month", month_code)
        .limit(1)
        .execute()
    )

    has_report = bool(res.data)
    status = res.data[0]["status"] if has_report else None
    submitted = has_report and status in ["Pending Review", "Submitted", "Approved"]

    return jsonify(
        {
            "has_report": has_report,
            "status": status,
            "submitted": submitted,
        }
    )

@pres.route(
    "/api/wallets/<int:wallet_id>/budgets/<int:budget_id>/reports/next-number",
    methods=["GET"],
)
def get_next_report_number(wallet_id, budget_id):
    """
    Helper para sa UI auto FR-00X, per wallet + budget (folder).
    """
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")
    try:
        res = (
            supabase.table("financial_reports")
            .select("report_no")
            .eq("organization_id", org_id)
            .eq("wallet_id", wallet_id)
            .eq("budget_id", budget_id)
            .order("created_at", desc=True)
            .limit(1)
            .execute()
        )
        if not res.data or not res.data[0].get("report_no"):
            return jsonify({"next_number": 1})

        last = res.data[0]["report_no"]
        try:
            num = int(last.split("-")[-1])
            return jsonify({"next_number": num + 1})
        except Exception:
            return jsonify({"next_number": 1})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@pres.route(
    "/api/wallets/<int:wallet_id>/budgets/<int:budget_id>/reports/latest",
    methods=["GET"],
)
def get_latest_report_for_budget(wallet_id, budget_id):
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")
    try:
        res = (
            supabase.table("financial_reports")
            .select(
                "id, status, report_no, event_name, date_prepared, "
                "budget, total_expense, reimbursement, previous_fund, budget_id"
            )
            .eq("organization_id", org_id)
            .eq("wallet_id", wallet_id)
            .eq("budget_id", budget_id)
            .order("created_at", desc=True)
            .limit(1)
            .execute()
        )
        if not res.data:
            return jsonify({"exists": False})
        return jsonify({"exists": True, "report": res.data[0]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# -----------------------
# Preview / Print (DOCX & HTML)
# -----------------------


@pres.route("/reports/<int:wallet_id>/budgets/<int:budget_id>/preview", methods=["GET"])
def preview_report_for_budget(wallet_id, budget_id):
    if not session.get("pres_user"):
        return "Unauthorized", 401

    org_id = session.get("org_id")

    rep_res = (
        supabase.table("financial_reports")
        .select("*")
        .eq("wallet_id", wallet_id)
        .eq("budget_id", budget_id)
        .eq("organization_id", org_id)
        .order("created_at", desc=True)
        .limit(1)
        .execute()
    )
    if not rep_res.data:
        return "No report", 404

    rep = rep_res.data[0]

    # org + college
    org_res = (
        supabase.table("organizations")
        .select("org_name, department_id")
        .eq("id", org_id)
        .single()
        .execute()
    )
    org_data = org_res.data or {}
    org_name = org_data.get("org_name") or ""

    college_name = "COLLEGE"
    dept_id = org_data.get("department_id")
    if dept_id is not None:
        dept_res = (
            supabase.table("departments")
            .select("dept_name")
            .eq("id", dept_id)
            .single()
            .execute()
        )
        if dept_res.data:
            college_name = dept_res.data["dept_name"].upper()

    # month text
    wb_res = (
        supabase.table("wallet_budgets")
        .select("year, month_id, months (month_name)")
        .eq("id", budget_id)
        .single()
        .execute()
    )
    report_month_text = ""
    if wb_res.data:
        y = wb_res.data["year"]
        mname = wb_res.data["months"]["month_name"]
        report_month_text = f"{mname} {y}".upper()

    # load template
    template_path = os.path.join(
        os.path.dirname(__file__),
        "templates",
        "pres",
        "finance_report_template.docx",
    )
    doc = Document(template_path)

    def replace_all(old: str, new: str):
        if new is None:
            new = ""
        for p in doc.paragraphs:
            if old in p.text:
                for run in p.runs:
                    run.text = run.text.replace(old, new)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old in cell.text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.text = run.text.replace(old, new)

    # numeric fields
    budget_val = float(rep.get("budget") or 0)
    total_expense = float(rep.get("total_expense") or 0)
    reimb = float(rep.get("reimbursement") or 0)
    prev_fund = float(rep.get("previous_fund") or 0)
    remaining = budget_val - total_expense - reimb + prev_fund
    total_income = float(rep.get("total_income") or 0)
    budget_in_the_bank = float(rep.get("budget_in_the_bank") or 0)

    # header placeholders
    replace_all("{{TOTAL_INCOME}}", f"PHP {total_income:,.2f}")
    replace_all("{{BUDGET_IN_THE_BANK}}", f"PHP {budget_in_the_bank:,.2f}")
    replace_all("{{COLLEGE_NAME}}", college_name)
    replace_all("{{ORG_NAME}}", org_name)
    replace_all("{{EVENT_NAME}}", rep.get("event_name") or "")
    replace_all("{{REPORT_MONTH}}", report_month_text)
    replace_all("{{DATE_PREPARED}}", str(rep.get("date_prepared") or ""))
    replace_all("{{REPORT_NO}}", rep.get("report_no") or "")
    replace_all("{{BUDGET}}", f"PHP {budget_val:,.2f}")
    replace_all("{{TOTAL_EXPENSE}}", f"PHP {total_expense:,.2f}")
    replace_all("{{REIMBURSEMENT}}", f"PHP {reimb:,.2f}")
    replace_all("{{PREVIOUS_FUND}}", f"PHP {prev_fund:,.2f}")
    replace_all("{{TOTAL_REMAINING}}", f"PHP {remaining:,.2f}")

    # INCOME rows
    income_res = (
        supabase.table("wallet_transactions")
        .select("date_issued, quantity, income_type, description, price, kind")
        .eq("wallet_id", wallet_id)
        .eq("budget_id", budget_id)
        .eq("kind", "income")
        .order("date_issued")
        .execute()
    )
    incomes = income_res.data or []

    income_table = None
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header_row = table.rows[1]
        header_text = " ".join(cell.text for cell in header_row.cells).upper()
        if "TYPE OF INCOME" in header_text and "DATE ISSUED" in header_text:
            income_table = table
            break

    # EXPENSE rows
    tx_res = (
        supabase.table("wallet_transactions")
        .select("date_issued, quantity, particulars, description, price, kind")
        .eq("wallet_id", wallet_id)
        .eq("budget_id", budget_id)
        .eq("kind", "expense")
        .order("date_issued")
        .execute()
    )
    txs = tx_res.data or []

    expenses_table = None
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header_row = table.rows[1]
        first_row_text = " ".join(cell.text for cell in header_row.cells).upper()
        if "DATE ISSUED" in first_row_text and "PARTICULARS" in first_row_text:
            expenses_table = table
            break

    # fill EXPENSES table
    if expenses_table:
        while len(expenses_table.rows) > 3:
            expenses_table._tbl.remove(expenses_table.rows[2]._tr)

        from datetime import datetime as _dt_local

        def fmt_date(d):
            try:
                return _dt_local.fromisoformat(d).strftime("%Y-%m-%d")
            except Exception:
                return d or ""

        last_date = None
        summary_row = expenses_table.rows[-1]

        for tx in txs:
            new_row = expenses_table.add_row()
            expenses_table._tbl.remove(new_row._tr)
            summary_row._tr.addprevious(new_row._tr)

            date_cell, qty_cell, part_cell, desc_cell, total_cell = new_row.cells

            date_str = tx.get("date_issued")
            show_date = fmt_date(date_str)
            if date_str == last_date:
                date_cell.text = ""
            else:
                date_cell.text = show_date
                last_date = date_str

            qty_cell.text = str(tx.get("quantity") or "")
            part_cell.text = tx.get("particulars") or ""
            desc_cell.text = tx.get("description") or ""

            qty = float(tx.get("quantity") or 0)
            price = float(tx.get("price") or 0)
            line_total = qty * price
            total_cell.text = f"PHP {line_total:,.2f}"

        summary_row.cells[-1].text = f"PHP {total_expense:,.2f}"

    # fill INCOME table
    if income_table:
        while len(income_table.rows) > 3:
            income_table._tbl.remove(income_table.rows[2]._tr)

        from datetime import datetime as _dt_local

        def fmt_date_income(d):
            try:
                return _dt_local.fromisoformat(d).strftime("%Y-%m-%d")
            except Exception:
                return d or ""

        last_date_inc = None
        summary_row_inc = income_table.rows[-1]

        for tx in incomes:
            new_row = income_table.add_row()
            income_table._tbl.remove(new_row._tr)
            summary_row_inc._tr.addprevious(new_row._tr)

            date_cell, qty_cell, type_cell, desc_cell, price_cell = new_row.cells

            date_str = tx.get("date_issued")
            show_date = fmt_date_income(date_str)
            if date_str == last_date_inc:
                date_cell.text = ""
            else:
                date_cell.text = show_date
                last_date_inc = date_str

            qty_cell.text = str(tx.get("quantity") or "")
            type_cell.text = tx.get("income_type") or ""
            desc_cell.text = tx.get("description") or ""

            price = float(tx.get("price") or 0)
            price_cell.text = f"PHP {price:,.2f}"

    # receipts -> pictures
    rc_res = (
        supabase.table("wallet_receipts")
        .select("description, file_url, receipt_date")
        .eq("wallet_id", wallet_id)
        .eq("budget_id", budget_id)
        .order("receipt_date")
        .execute()
    )
    receipts = rc_res.data or []

    if receipts:
        # walang page break, diretso lang sa current page
        for r in receipts:
            file_path = r["file_url"]
            try:
                file_bytes = supabase.storage.from_(BUCKET_RECEIPTS).download(file_path)
            except Exception:
                continue

            # caption paragraph (center)
            cap_p = doc.add_paragraph()
            cap_p.alignment = 1  # 0=left, 1=center, 2=right
            cap_p.add_run(f"{r['receipt_date']} - {r['description']}")

            # image paragraph (center)
            img_stream = BytesIO(file_bytes)
            doc.add_picture(img_stream, width=Inches(3))
            last_par = doc.paragraphs[-1]
            last_par.alignment = 1

            # maliit na spacer sa ilalim ng bawat picture
            spacer = doc.add_paragraph()
            spacer.alignment = 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=False,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        download_name="financial_report_preview.docx",
    )


@pres.route("/reports/<int:wallet_id>/budgets/<int:budget_id>/print", methods=["GET"])
def print_report_for_budget(wallet_id, budget_id):
    if not session.get("pres_user"):
        return redirect(url_for("pres.pres_login"))

    org_id = session.get("org_id")

    rep_res = (
        supabase.table("financial_reports")
        .select("*")
        .eq("wallet_id", wallet_id)
        .eq("budget_id", budget_id)
        .eq("organization_id", org_id)
        .order("created_at", desc=True)
        .limit(1)
        .execute()
    )
    if not rep_res.data:
        return "No report", 404

    rep = rep_res.data[0]

    # org + college
    org_res = (
        supabase.table("organizations")
        .select("org_name, department_id")
        .eq("id", org_id)
        .single()
        .execute()
    )
    org_data = org_res.data or {}
    org_name = org_data.get("org_name") or ""

    college_name = "COLLEGE"
    dept_id = org_data.get("department_id")
    if dept_id is not None:
        dept_res = (
            supabase.table("departments")
            .select("dept_name")
            .eq("id", dept_id)
            .single()
            .execute()
        )
        if dept_res.data:
            college_name = dept_res.data["dept_name"].upper()

    # month text
    wb_res = (
        supabase.table("wallet_budgets")
        .select("year, month_id, months (month_name)")
        .eq("id", budget_id)
        .single()
        .execute()
    )
    report_month_text = ""
    if wb_res.data:
        y = wb_res.data["year"]
        mname = wb_res.data["months"]["month_name"]
        report_month_text = f"{mname} {y}".upper()

    # numeric fields
    budget_val = float(rep.get("budget") or 0)
    total_expense = float(rep.get("total_expense") or 0)
    reimb = float(rep.get("reimbursement") or 0)
    prev_fund = float(rep.get("previous_fund") or 0)
    remaining = budget_val - total_expense - reimb + prev_fund
    total_income = float(rep.get("total_income") or 0)
    budget_in_the_bank = float(rep.get("budget_in_the_bank") or 0)

    # expenses for table sa print_report.html
    tx_res = (
        supabase.table("wallet_transactions")
        .select("date_issued, quantity, particulars, description, price, kind")
        .eq("wallet_id", wallet_id)
        .eq("budget_id", budget_id)
        .eq("kind", "expense")
        .order("date_issued")
        .execute()
    )
    txs = tx_res.data or []

    # Income transactions
    inc_res = (
        supabase.table("wallet_transactions")
        .select("date_issued, quantity, income_type, description, price, kind")
        .eq("wallet_id", wallet_id)
        .eq("budget_id", budget_id)
        .eq("kind", "income")
        .order("date_issued")
        .execute()
    )
    incomes = inc_res.data or []
    # receipts with public URL
    rc_res = (
        supabase.table("wallet_receipts")
        .select("description, file_url, receipt_date")
        .eq("wallet_id", wallet_id)
        .eq("budget_id", budget_id)
        .order("receipt_date")
        .execute()
    )
    receipts = rc_res.data or []
    for r in receipts:
        try:
            public_url = supabase.storage.from_(BUCKET_RECEIPTS).get_public_url(
                r["file_url"]
            )
            if isinstance(public_url, dict):
                r["file_url"] = public_url.get("publicUrl") or public_url.get(
                    "signedURL"
                )
            else:
                r["file_url"] = public_url
        except Exception:
            r["file_url"] = ""

    return render_template(
        "print_report.html",
        report=rep,
        budget=budget_val,
        totalexpense=total_expense,
        reimbursement=reimb,
        previous_fund=prev_fund,
        remaining=remaining,
        total_income=total_income,
        budget_in_the_bank=budget_in_the_bank,
        transactions=txs,
        incomes=incomes,
        org_name=org_name,
        college_name=college_name,
        report_month_text=report_month_text,
        receipts=receipts,
    )


# -----------------------
# Submit report -> archive + reset
# -----------------------
from datetime import datetime as dt  # siguraduhin na nasa taas na ito


def update_osas_financial_report(org_id, budget_id):
    """
    Update OSAS financial_reports row for this org based on submitted month.
    Marks the month as received in checklist and recomputes status.
    """
    # Kunin month name mula sa wallet_budgets + months table
    wb_res = (
        supabase.table("wallet_budgets")
        .select("months (month_name)")
        .eq("id", budget_id)
        .single()
        .execute()
    )
    if not wb_res.data:
        return

    month_name = (wb_res.data["months"]["month_name"] or "").lower()
    month_key = month_name  # e.g. "january", "february", ...

    # Hanapin OSAS financial_reports row ng org
    fr_res = (
        supabase.table("financial_reports")
        .select("*")
        .eq("organization_id", org_id)
        .is_("wallet_id", None)  # master row lang
        .is_("budget_id", None)
        .limit(1)
        .execute()
    )

    report = fr_res.data[0]
    checklist = report.get("checklist") or {}

    month_keys = [
        "august",
        "september",
        "october",
        "november",
        "december",
        "january",
        "february",
        "march",
        "april",
        "may",
    ]

    if month_key in month_keys:
        checklist[month_key] = True

    received_count = sum(1 for k in month_keys if checklist.get(k))
    total_count = len(month_keys)

    if received_count == 0:
        status = "Pending Review"
    elif received_count < total_count:
        status = "In Review"
    else:
        status = "Completed"

    supabase.table("financial_reports").update(
        {
            "checklist": checklist,
            "status": status,
            "updated_at": dt.utcnow().isoformat(),
        }
    ).eq("id", report["id"]).execute()


@pres.route(
    "/api/wallets/<int:wallet_id>/budgets/<int:budget_id>/submit",
    methods=["GET"],
)
def check_report_submitted(wallet_id, budget_id):
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")
    try:
        res = (
            supabase.table("financial_reports")
            .select("status")
            .eq("organization_id", org_id)
            .eq("wallet_id", wallet_id)
            .eq("budget_id", budget_id)
            .order("created_at", desc=True)
            .limit(1)
            .execute()
        )
        if not res.data:
            return jsonify({"submitted": False})
        submitted = res.data[0].get("status") == "Submitted"
        return jsonify({"submitted": submitted})
    except Exception as e:
        return jsonify({"error": str(e), "submitted": False}), 500


@pres.route("/reports/<int:wallet_id>/submit", methods=["POST"])
def submitreportwalletid(wallet_id):
    """
    Mark latest Pending report as Submitted, create archive snapshot
    (summary + transactions + receipts), then clear month data and reset budget.
    JS calls: POST /pres/reports/<walletId>/submit
    """
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")

    try:
        repres = (
            supabase.table("financial_reports")
            .select("*")
            .eq("organization_id", org_id)
            .eq("wallet_id", wallet_id)
            .eq("status", "Pending Review")
            .order("created_at", desc=True)
            .limit(1)
            .execute()
        )
        if not repres.data:
            return jsonify({"error": "No pending report"}), 404

        rep = repres.data[0]
        rep_id = rep["id"]
        budget_id = rep["budget_id"]

        # Mark report as submitted
        supabase.table("financial_reports").update(
            {
                "status": "Submitted",
                "submission_date": dt.utcnow().date().isoformat(),
                "updated_at": dt.utcnow().isoformat(),
            }
        ).eq("id", rep_id).execute()
        create_osas_notification(org_id, rep_id, 'has a report "Pending Review"')

        # Compute remaining balance
        budget_val = float(rep.get("budget") or 0)
        total_expense = float(rep.get("total_expense") or 0)
        reimb = float(rep.get("reimbursement") or 0)
        prev_fund = float(rep.get("previous_fund") or 0)
        remaining = budget_val - total_expense - reimb + prev_fund

        # Insert archive summary
        arch_ins = (
            supabase.table("financial_report_archives")
            .insert(
                {
                    "organization_id": org_id,
                    "wallet_id": wallet_id,
                    "budget_id": budget_id,
                    "report_id": rep_id,
                    "report_no": rep.get("report_no"),
                    "event_name": rep.get("event_name"),
                    "date_prepared": rep.get("date_prepared"),
                    "budget": budget_val,
                    "total_expense": total_expense,
                    "reimbursement": reimb,
                    "previous_fund": prev_fund,
                    "remaining": remaining,
                    "file_url": None,
                }
            )
            .execute()
        )
        archive_id = arch_ins.data[0]["id"]

        # Archive transactions
        txres = (
            supabase.table("wallet_transactions")
            .select("date_issued, quantity, particulars, description, price, kind")
            .eq("wallet_id", wallet_id)
            .eq("budget_id", budget_id)
            .execute()
        )
        txs = txres.data or []
        if txs:
            tx_rows = [
                {
                    "archive_id": archive_id,
                    "date_issued": tx["date_issued"],
                    "quantity": tx["quantity"],
                    "particulars": tx.get("particulars"),
                    "description": tx["description"],
                    "price": float(tx["price"]),
                    "kind": tx["kind"],
                }
                for tx in txs
            ]
            supabase.table("financial_report_archive_transactions").insert(
                tx_rows
            ).execute()

        # Archive receipts
        rcres = (
            supabase.table("wallet_receipts")
            .select("description, receipt_date, file_url")
            .eq("wallet_id", wallet_id)
            .eq("budget_id", budget_id)
            .execute()
        )
        receipts = rcres.data or []
        if receipts:
            rc_rows = [
                {
                    "archive_id": archive_id,
                    "description": r["description"],
                    "receipt_date": r["receipt_date"],
                    "file_url": r["file_url"],
                }
                for r in receipts
            ]
            supabase.table("financial_report_archive_receipts").insert(
                rc_rows
            ).execute()

        # Clear current month data
        supabase.table("wallet_transactions").delete().eq("wallet_id", wallet_id).eq(
            "budget_id", budget_id
        ).execute()

        supabase.table("wallet_receipts").delete().eq("wallet_id", wallet_id).eq(
            "budget_id", budget_id
        ).execute()

        supabase.table("wallet_budgets").update({"amount": 0}).eq(
            "id", budget_id
        ).execute()

        # Update OSAS-side financial_reports checklist/status
        update_osas_financial_report(org_id, budget_id)

        return jsonify({"success": True, "archive_id": archive_id}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# -----------------------
# Archives (submitted reports per folder)
# -----------------------


@pres.route("/api/wallets/<int:folder_id>/archives", methods=["GET"])
def get_wallet_archives(folder_id):
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")
    wallet_id = get_real_wallet_id(folder_id)
    if wallet_id is None:
        return jsonify({"error": "Wallet folder not found"}), 404

    try:
        res = (
            supabase.table("financial_report_archives")
            .select(
                "id, report_id, report_no, event_name, date_prepared, "
                "budget, total_expense, reimbursement, previous_fund, remaining, file_url"
            )
            .eq("organization_id", org_id)
            .eq("wallet_id", wallet_id)
            .eq("budget_id", folder_id)
            .order("created_at")
            .execute()
        )

        # ✅ EXACT field names matching JS!
        archives = [
            {
                "id": a["id"],
                "reportid": a["report_id"],  # ← lowercase reportid
                "reportno": a["report_no"],  # ← lowercase reportno
                "eventname": a["event_name"],  # ← lowercase eventname
                "dateprepared": a["date_prepared"],  # ← lowercase dateprepared
                "budget": float(a["budget"] or 0),
                "totalexpense": float(
                    a["total_expense"] or 0
                ),  # ← lowercase totalexpense
                "reimbursement": float(a["reimbursement"] or 0),
                "previousfund": float(
                    a["previous_fund"] or 0
                ),  # ← lowercase previousfund
                "remaining": float(a["remaining"] or 0),
                "fileurl": a["file_url"],  # ← lowercase fileurl
            }
            for a in (res.data or [])
        ]
        return jsonify(archives)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


from io import BytesIO
from docx import Document
from docx.shared import Inches
from datetime import datetime as _dt


@pres.route("/api/archives/<int:archive_id>/download", methods=["GET"])
def download_archive(archive_id):
    """
    Generate a DOCX file for a submitted (archived) report on the fly and return it.
    Ginagamit ang financial_report_archives + archive_transactions + archive_receipts.
    Final URL: /pres/api/archives/<archive_id>/download
    """
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")

    try:
        # 1) Kunin archive summary row
        arch_res = (
            supabase.table("financial_report_archives")
            .select("*")
            .eq("id", archive_id)
            .single()
            .execute()
        )
        if not arch_res.data:
            return jsonify({"error": "Archive not found"}), 404

        arch = arch_res.data
        wallet_id = arch["wallet_id"]
        budget_id = arch["budget_id"]

        # 2) org + college name
        org_res = (
            supabase.table("organizations")
            .select("org_name, department_id")
            .eq("id", org_id)
            .single()
            .execute()
        )
        org_data = org_res.data or {}
        org_name = org_data.get("org_name") or ""

        college_name = "COLLEGE"
        dept_id = org_data.get("department_id")
        if dept_id is not None:
            dept_res = (
                supabase.table("departments")
                .select("dept_name")
                .eq("id", dept_id)
                .single()
                .execute()
            )
            if dept_res.data:
                college_name = dept_res.data["dept_name"].upper()

        # 3) month info ng wallet/budget (REPORT_MONTH)
        wb_res = (
            supabase.table("wallet_budgets")
            .select("year, month_id, months (month_name)")
            .eq("id", budget_id)
            .single()
            .execute()
        )
        report_month_text = ""
        if wb_res.data:
            y = wb_res.data["year"]
            mname = wb_res.data["months"]["month_name"]
            report_month_text = f"{mname} {y}".upper()

        # 4) load template
        template_path = os.path.join(
            os.path.dirname(__file__),
            "templates",
            "pres",
            "finance_report_template.docx",
        )
        doc = Document(template_path)

        def replace_all(old: str, new: str):
            if new is None:
                new = ""
            for p in doc.paragraphs:
                if old in p.text:
                    for run in p.runs:
                        run.text = run.text.replace(old, new)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old in cell.text:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.text = run.text.replace(old, new)

        # 5) numeric values from archive row
        budget_val = float(arch.get("budget") or 0)
        total_expense = float(arch.get("total_expense") or 0)
        reimb = float(arch.get("reimbursement") or 0)
        prev_fund = float(arch.get("previous_fund") or 0)
        remaining = float(arch.get("remaining") or 0)

        # 6) header / summary placeholders
        replace_all("{{COLLEGE_NAME}}", college_name)
        replace_all("{{ORG_NAME}}", org_name)
        replace_all("{{EVENT_NAME}}", arch.get("event_name") or "")
        replace_all("{{REPORT_MONTH}}", report_month_text)
        replace_all("{{DATE_PREPARED}}", str(arch.get("date_prepared") or ""))
        replace_all("{{REPORT_NO}}", arch.get("report_no") or "")
        replace_all("{{BUDGET}}", f"PHP {budget_val:,.2f}")
        replace_all("{{TOTAL_EXPENSE}}", f"PHP {total_expense:,.2f}")
        replace_all("{{REIMBURSEMENT}}", f"PHP {reimb:,.2f}")
        replace_all("{{PREVIOUS_FUND}}", f"PHP {prev_fund:,.2f}")
        replace_all("{{TOTAL_REMAINING}}", f"PHP {remaining:,.2f}")

        # 7) expenses table rows galing sa financial_report_archive_transactions
        tx_res = (
            supabase.table("financial_report_archive_transactions")
            .select("date_issued, quantity, particulars, description, price, kind")
            .eq("archive_id", archive_id)
            .eq("kind", "expense")
            .order("date_issued")
            .execute()
        )
        txs = tx_res.data or []

        expenses_table = None
        for table in doc.tables:
            if len(table.rows) < 2:
                continue
            header_row = table.rows[1]
            first_row_text = " ".join(cell.text for cell in header_row.cells).upper()
            if "DATE ISSUED" in first_row_text and "PARTICULARS" in first_row_text:
                expenses_table = table
                break

        if expenses_table:
            while len(expenses_table.rows) > 3:
                expenses_table._tbl.remove(expenses_table.rows[2]._tr)

            def fmt_date(d):
                try:
                    return _dt.fromisoformat(d).strftime("%Y-%m-%d")
                except Exception:
                    return d or ""

            last_date = None
            summary_row = expenses_table.rows[-1]

            for tx in txs:
                new_row = expenses_table.add_row()
                expenses_table._tbl.remove(new_row._tr)
                summary_row._tr.addprevious(new_row._tr)

                date_cell, qty_cell, part_cell, desc_cell, total_cell = new_row.cells

                date_str = tx.get("date_issued")
                show_date = fmt_date(date_str)
                if date_str == last_date:
                    date_cell.text = ""
                else:
                    date_cell.text = show_date
                    last_date = date_str

                qty_cell.text = str(tx.get("quantity") or "")
                part_cell.text = tx.get("particulars") or ""
                desc_cell.text = tx.get("description") or ""

                qty = float(tx.get("quantity") or 0)
                price = float(tx.get("price") or 0)
                line_total = qty * price
                total_cell.text = f"PHP {line_total:,.2f}"

            summary_row.cells[-1].text = f"PHP {total_expense:,.2f}"

        # 8) receipts appendix galing sa financial_report_archive_receipts
        rc_res = (
            supabase.table("financial_report_archive_receipts")
            .select("description, file_url, receipt_date")
            .eq("archive_id", archive_id)
            .order("receipt_date")
            .execute()
        )
        receipts = rc_res.data or []

        if receipts:
            doc.add_page_break()
            title_p = doc.add_paragraph()
            title_run = title_p.add_run("APPENDIX: RECEIPTS")
            title_run.bold = True

            for r in receipts:
                file_path = r["file_url"]
                try:
                    file_bytes = supabase.storage.from_(BUCKET_RECEIPTS).download(
                        file_path
                    )
                except Exception:
                    continue

                cap_p = doc.add_paragraph()
                cap_p.add_run(f"{r['receipt_date']} - {r['description']}")

                img_stream = BytesIO(file_bytes)
                doc.add_picture(img_stream, width=Inches(4))
                doc.add_paragraph()

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        download_name = arch.get("report_no") or "financial_report.docx"
        return send_file(
            buf,
            as_attachment=True,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            download_name=download_name,
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# -----------------------
# Profile APIs
# -----------------------
BUCKET_PROFILE_PHOTO = "profile_photo"  # supabase bucket name


@pres.route("/api/profile", methods=["GET"])
def get_profile():
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")

    try:
        # main org row (OSAS table)
        org_res = (
            supabase.table("organizations")
            .select(
                "id, org_name, department_id, "
                "accreditation_date, status"
            )
            .eq("id", org_id)
            .limit(1)
            .execute()
        )
        if not org_res.data:
            return jsonify({"error": "Organization not found"}), 404

        org = org_res.data[0]

        # PRES-only profile row
        prof_res = (
            supabase.table("profile_users")
            .select("org_short_name, campus, school_name, profile_photo_url")
            .eq("organization_id", org_id)
            .limit(1)
            .execute()
        )
        prof = prof_res.data[0] if prof_res.data else {}

        # department name
        dept_name = None
        if org.get("department_id") is not None:
            dept_res = (
                supabase.table("departments")
                .select("dept_name")
                .eq("id", org["department_id"])
                .limit(1)
                .execute()
            )
            if dept_res.data:
                dept_name = dept_res.data[0]["dept_name"]

        # optional public URL for profile photo (from profile_users)
        photo_url = None
        if prof.get("profile_photo_url"):
            public = supabase.storage.from_(BUCKET_PROFILE_PHOTO).get_public_url(
                prof["profile_photo_url"]
            )
            photo_url = public.get("publicUrl") if isinstance(public, dict) else public

        # accreditation info purely from OSAS organizations table
        status_text = "Accredited" if org.get("status") == "Active" else org.get("status")

        profile = {
            "org_name": org.get("org_name"),
            "org_short_name": prof.get("org_short_name"),
            "department": dept_name,
            "department_id": org.get("department_id"),
            "school": prof.get("school_name"),
            "profile_photo_url": photo_url,
            "accreditation": {
                "date_of_accreditation": org.get("accreditation_date"),
                "current_status": status_text,
            },
        }

        return jsonify(profile)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@pres.route("/api/profile", methods=["PUT", "POST"])
def update_profile():
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")
    data = request.get_json() or {}

    try:
        org_update = {}
        prof_update = {}

        # OSAS org table fields
        if "org_name" in data:
            org_update["org_name"] = data["org_name"]
        if "department_id" in data:
            org_update["department_id"] = data["department_id"]

        # PRES profile fields (profile_users table)
        if "org_short_name" in data:
            prof_update["org_short_name"] = data["org_short_name"]
        if "school" in data:
            prof_update["school_name"] = data["school"]
        if "campus" in data:
            prof_update["campus"] = data["campus"]

        # Update organizations
        if org_update:
            supabase.table("organizations").update(org_update).eq(
                "id", org_id
            ).execute()
            if "org_name" in data:
                session["org_name"] = data["org_name"]

               # Upsert profile_users
        if prof_update:
            existing = (
                supabase.table("profile_users")
                .select("id")
                .eq("organization_id", org_id)
                .limit(1)
                .execute()
            )

            try:
                if existing.data:
                    supabase.table("profile_users").update(prof_update).eq(
                        "organization_id", org_id
                    ).execute()
                else:
                    prof_update["organization_id"] = org_id
                    supabase.table("profile_users").insert(prof_update).execute()
            except Exception as e:
                # duplicate short name (unique violation)
                msg = str(e)
                if "duplicate key value violates unique constraint" in msg or "23505" in msg:
                    return (
                        jsonify(
                            {
                                "success": False,
                                "error": "Shortened name is already used by another organization.",
                            }
                        ),
                        400,
                    )
                return jsonify({"success": False, "error": "Failed to update profile."}), 500

        if org_update or prof_update:
            return jsonify(
                {"success": True, "message": "Profile updated successfully"}
            )

        return jsonify({"error": "No valid fields to update"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@pres.route("/api/profile/picture", methods=["POST"])
def upload_profile_picture():
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")

    if "photo" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["photo"]

    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    allowed_extensions = {"png", "jpg", "jpeg", "gif"}
    file_ext = file.filename.rsplit(".", 1)[1].lower() if "." in file.filename else ""
    if file_ext not in allowed_extensions:
        return jsonify({"error": "Invalid file type"}), 400

    try:
        from uuid import uuid4

        raw_bytes = file.read()
        filename = f"org-{org_id}-{uuid4()}.{file_ext}"
        path = f"orgs/{org_id}/{filename}"

        # upload to Supabase Storage
        supabase.storage.from_(BUCKET_PROFILE_PHOTO).upload(path, raw_bytes)

        # upsert into profile_users.profile_photo_url
        existing = (
            supabase.table("profile_users")
            .select("id")
            .eq("organization_id", org_id)
            .limit(1)
            .execute()
        )
        if existing.data:
            supabase.table("profile_users").update(
                {"profile_photo_url": path}
            ).eq("organization_id", org_id).execute()
        else:
            supabase.table("profile_users").insert(
                {
                    "organization_id": org_id,
                    "profile_photo_url": path,
                }
            ).execute()

        # return public URL
        public = supabase.storage.from_(BUCKET_PROFILE_PHOTO).get_public_url(path)
        url = public.get("publicUrl") if isinstance(public, dict) else public

        return jsonify(
            {
                "success": True,
                "message": "Profile picture uploaded successfully",
                "path": path,
                "url": url,
            }
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# -----------------------
# Officers APIs
# -----------------------


@pres.route("/api/officers", methods=["GET"])
def get_officers():
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")
    res = (
        supabase.table("profile_officers")
        .select("id, name, position, term_start, term_end, status")
        .eq("organization_id", org_id)
        .order("term_start")
        .execute()
    )
    return jsonify({"officers": res.data or []})


@pres.route("/api/officers", methods=["POST"])
def create_officer():
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")
    data = request.get_json() or {}
    officer = {
        "organization_id": org_id,
        "name": data.get("name"),
        "position": data.get("position"),
        "term_start": data.get("term_start"),
        "term_end": data.get("term_end"),
        "status": data.get("status") or "Active",
    }
    res = supabase.table("profile_officers").insert(officer).execute()
    created = res.data[0]
    return jsonify({"officer": created}), 201


@pres.route("/api/officers/<int:officer_id>", methods=["PUT"])
def update_officer(officer_id):
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")
    data = request.get_json() or {}
    update = {
        "name": data.get("name"),
        "position": data.get("position"),
        "term_start": data.get("term_start"),
        "term_end": data.get("term_end"),
        "status": data.get("status"),
    }
    update = {k: v for k, v in update.items() if v is not None}

    res = (
        supabase.table("profile_officers")
        .update(update)
        .eq("id", officer_id)
        .eq("organization_id", org_id)
        .execute()
    )
    if not res.data:
        return jsonify({"error": "Not found"}), 404
    return jsonify({"officer": res.data[0]})


@pres.route("/api/officers/<int:officer_id>", methods=["DELETE"])
def delete_officer(officer_id):
    if not session.get("pres_user"):
        return jsonify({"error": "Unauthorized"}), 401

    org_id = session.get("org_id")
    supabase.table("profile_officers").delete().eq("id", officer_id).eq(
        "organization_id", org_id
    ).execute()
    return jsonify({"success": True})